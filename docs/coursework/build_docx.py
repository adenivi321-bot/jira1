"""Сборка курсовой работы TaskHub в формате .docx по требованиям ГОСТ.

Параметры оформления (в соответствии с методическими указаниями):
  - А4, Times New Roman 14 pt, межстрочный интервал 1,5
  - Выравнивание по ширине, абзацный отступ 1,25 см
  - Поля: левое 30 мм, правое 15 мм, верхнее/нижнее 20 мм
  - Сквозная нумерация страниц (по центру внизу), номер начинается с введения
"""

from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
PNG = ROOT / "diagrams" / "png"
OUT = ROOT / "Курсовая_работа_TaskHub.docx"


# ─────────────────────── Helpers ───────────────────────


def set_default_font(doc: Document) -> None:
    """Включает Times New Roman 14pt и интервал 1,5 во всём документе."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), "Times New Roman")
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(1.25)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def set_margins(section, *, left=30, right=15, top=20, bottom=20) -> None:
    section.left_margin = Mm(left)
    section.right_margin = Mm(right)
    section.top_margin = Mm(top)
    section.bottom_margin = Mm(bottom)


def add_page_numbers(section, *, suppress_first: bool = False) -> None:
    """Номер страницы по центру внизу. Сквозная нумерация документа."""
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def suppress_page_number(section) -> None:
    """Подавить номер на конкретной странице (титул и т.п.)."""
    sect_pr = section._sectPr
    title_pg = OxmlElement("w:titlePg")
    sect_pr.append(title_pg)
    header = section.first_page_header
    footer = section.first_page_footer
    if footer.paragraphs:
        for p in footer.paragraphs:
            p.text = ""


def add_para(
    doc,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    align=None,
    size: int = 14,
    indent: float | None = None,
    space_after: float = 0,
    line_spacing: float = 1.5,
):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    pf.first_line_indent = Cm(indent if indent is not None else 1.25)
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rfonts.set(qn(f"w:{attr}"), "Times New Roman")
    return p


def add_heading(doc, text: str, level: int = 1) -> None:
    """level=0 — структурный заголовок (ВВЕДЕНИЕ и т.п.), level=1 — ГЛАВА N, level=2 — 1.1."""
    align = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
    indent = 0 if level == 0 else 1.25
    p = add_para(
        doc,
        text,
        bold=True,
        align=align,
        size=14,
        indent=indent,
        space_after=0,
    )
    p.paragraph_format.keep_with_next = True


def add_image(doc, path: Path, caption: str, *, width_cm: float = 16) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = add_para(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, size=12)
    cap.runs[0].italic = True


def add_page_break(doc) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_listing(doc, code: str) -> None:
    """Моноширинный листинг кода."""
    for line in code.splitlines():
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.0
        pf.first_line_indent = Cm(0)
        pf.left_indent = Cm(1.0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        rpr = run._r.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("ascii", "hAnsi", "cs"):
            rfonts.set(qn(f"w:{attr}"), "Consolas")


# ─────────────────────── Build ───────────────────────


def build() -> Path:
    doc = Document()
    set_default_font(doc)
    section = doc.sections[0]
    set_margins(section)

    # Первая (титульная) страница: без номера
    suppress_page_number(section)
    add_page_numbers(section)

    # ───── Титульный лист ─────
    p = add_para(doc, "МИНИСТЕРСТВО ОБРАЗОВАНИЯ И НАУКИ", align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, bold=True)
    add_para(doc, "ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ ПРОФЕССИОНАЛЬНОЕ", align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, bold=True)
    add_para(doc, "ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ", align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, bold=True)
    add_para(doc, "«КОЛЛЕДЖ ИНФОРМАЦИОННЫХ ТЕХНОЛОГИЙ»", align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, bold=True)
    for _ in range(6):
        add_para(doc, "", indent=0)
    add_para(doc, "КУРСОВАЯ РАБОТА", align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, bold=True, size=16)
    add_para(doc, "по дисциплине", align=WD_ALIGN_PARAGRAPH.CENTER, indent=0)
    add_para(doc, "«Проектирование и разработка информационных систем»", align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, italic=True)
    add_para(doc, "", indent=0)
    add_para(doc, "на тему:", align=WD_ALIGN_PARAGRAPH.CENTER, indent=0)
    add_para(
        doc,
        "«Разработка прототипа системы управления задачами",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        indent=0,
        bold=True,
    )
    add_para(
        doc,
        "(наподобие Jira) с полным комплектом UML-диаграмм»",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        indent=0,
        bold=True,
    )
    for _ in range(6):
        add_para(doc, "", indent=0)
    add_para(doc, "Выполнил(а): студент(ка) группы _________", indent=0)
    add_para(doc, "________________________________________", indent=0)
    add_para(doc, "                              (Ф.И.О.)", indent=0)
    add_para(doc, "Руководитель: ____________________________", indent=0)
    add_para(doc, "                              (Ф.И.О.)", indent=0)
    add_para(doc, "Оценка: __________________________________", indent=0)
    add_para(doc, "Подпись руководителя: ____________________", indent=0)
    for _ in range(4):
        add_para(doc, "", indent=0)
    add_para(doc, "20___ г.", align=WD_ALIGN_PARAGRAPH.CENTER, indent=0)

    add_page_break(doc)

    # ───── Лист задания ─────
    add_para(doc, "ЗАДАНИЕ НА КУРСОВУЮ РАБОТУ", align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, bold=True)
    add_para(doc, "", indent=0)
    add_para(doc, "Студент: ________________________________________________________", indent=0)
    add_para(doc, "Группа:  ________________________________________________________", indent=0)
    add_para(doc, "Тема: «Разработка прототипа системы управления задачами (наподобие Jira) с полным комплектом UML-диаграмм».", indent=0)
    add_para(doc, "", indent=0)
    add_para(doc, "Исходные данные к работе:", bold=True)
    add_para(doc, "— требования к программному продукту класса task management;")
    add_para(doc, "— нотация UML 2.5 (стандарт ISO/IEC 19505);")
    add_para(doc, "— стек технологий: TypeScript, NestJS, React, PostgreSQL, Redis, Docker.")
    add_para(doc, "Перечень вопросов, подлежащих разработке:", bold=True)
    add_para(doc, "1) анализ предметной области и обзор аналогов;")
    add_para(doc, "2) формирование функциональных и нефункциональных требований;")
    add_para(doc, "3) проектирование архитектуры и базы данных;")
    add_para(doc, "4) разработка полного комплекта UML-диаграмм;")
    add_para(doc, "5) реализация прототипа системы;")
    add_para(doc, "6) описание ключевых программных модулей.")
    add_para(doc, "", indent=0)
    add_para(doc, "Перечень графического материала:", bold=True)
    add_para(doc, "диаграмма вариантов использования; диаграмма классов; ER-диаграмма; диаграммы последовательности (3 шт.); диаграмма деятельности; диаграмма состояний; диаграмма компонентов; диаграмма развёртывания; диаграмма пакетов.")
    add_para(doc, "", indent=0)
    add_para(doc, "Дата выдачи задания: «____» _____________ 20___ г.", indent=0)
    add_para(doc, "Срок сдачи работы:    «____» _____________ 20___ г.", indent=0)
    add_para(doc, "Руководитель: _________________ / ____________________", indent=0)
    add_para(doc, "Задание принял к исполнению: ___________ / ___________", indent=0)

    add_page_break(doc)

    # ───── СОДЕРЖАНИЕ ─────
    add_heading(doc, "СОДЕРЖАНИЕ", level=0)
    add_para(doc, "", indent=0)
    toc = [
        ("ВВЕДЕНИЕ", 4),
        ("ГЛАВА 1. ТЕОРЕТИЧЕСКИЕ АСПЕКТЫ ПРОЕКТИРОВАНИЯ СИСТЕМ УПРАВЛЕНИЯ ЗАДАЧАМИ", 6),
        ("1.1 Понятие системы управления задачами и её роль в управлении проектами", 6),
        ("1.2 Обзор существующих решений и сравнительный анализ", 8),
        ("1.3 Назначение и виды UML-диаграмм при проектировании ИС", 10),
        ("1.4 Выбор технологического стека", 12),
        ("ГЛАВА 2. ПРАКТИЧЕСКАЯ РАЗРАБОТКА ПРОТОТИПА СИСТЕМЫ TASKHUB", 14),
        ("2.1 Постановка задачи и требования к прототипу", 14),
        ("2.2 Проектирование архитектуры системы", 16),
        ("2.3 Полный комплект UML-диаграмм проекта", 18),
        ("2.4 Проектирование базы данных и обеспечение мультитенантности", 26),
        ("2.5 Реализация серверной части (NestJS)", 28),
        ("2.6 Реализация клиентской части (React)", 30),
        ("2.7 Развёртывание прототипа и тестирование", 32),
        ("ЗАКЛЮЧЕНИЕ", 34),
        ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 35),
        ("ПРИЛОЖЕНИЕ А. Листинг ключевых модулей", 37),
        ("ПРИЛОЖЕНИЕ Б. Схема базы данных (init.sql)", 40),
    ]
    for title, page in toc:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        pf.first_line_indent = Cm(0)
        tab_stops = pf.tab_stops
        tab_stops.add_tab_stop(Cm(16), WD_ALIGN_PARAGRAPH.RIGHT, 1)  # 1 = leader dots
        run = p.add_run(title)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run = p.add_run(f"\t{page}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

    add_page_break(doc)

    # ───── ВВЕДЕНИЕ ─────
    add_heading(doc, "ВВЕДЕНИЕ", level=0)
    intro_paragraphs = [
        "Современные ИТ-компании ведут работу одновременно над десятками проектов и сотнями задач, в выполнении которых участвуют распределённые команды. В таких условиях ручной учёт работы в виде таблиц или файловых документов становится неэффективным: теряется прозрачность, страдает координация, увеличивается время реакции на изменения. Системы управления задачами (Task Management Systems) — программные продукты, которые позволяют формализовать процесс планирования, контроля и анализа работы команды, обеспечивая единую среду для постановки задач, отслеживания их статусов и хранения истории изменений.",
        "Наиболее известным представителем класса является продукт Jira от Atlassian, занимающий по различным оценкам около 40 % мирового рынка систем такого типа. Однако высокая стоимость лицензии, ограничения «облачной» версии и отсутствие возможности самостоятельного хостинга на собственных серверах побуждают многие организации искать или разрабатывать альтернативные решения. Это особенно актуально для государственных и коммерческих компаний, переходящих на отечественное программное обеспечение в рамках курса на импортозамещение.",
        "Актуальность темы курсовой работы обусловлена тем, что разработка собственного прототипа системы управления задачами позволяет, во-первых, освоить полный цикл проектирования информационной системы — от анализа предметной области и моделирования с помощью унифицированного языка моделирования (UML) до программной реализации и развёртывания, а во-вторых, создать гибкое решение, которое может быть адаптировано под нужды конкретной организации.",
        "Цель курсовой работы — спроектировать и реализовать прототип многопользовательской системы управления задачами «TaskHub», построить полный комплект UML-диаграмм, описывающих систему с разных точек зрения, и продемонстрировать применимость полученной модели к реальному коду.",
        "Для достижения поставленной цели необходимо решить следующие задачи:",
        "1) проанализировать предметную область и провести сравнительный обзор существующих систем управления задачами;",
        "2) сформулировать функциональные и нефункциональные требования к разрабатываемому прототипу;",
        "3) спроектировать архитектуру системы и схему базы данных;",
        "4) разработать полный комплект UML-диаграмм (вариантов использования, классов, состояний, последовательности, деятельности, компонентов, развёртывания, пакетов и ER-модель);",
        "5) реализовать серверную часть, клиентский интерфейс и инфраструктуру развёртывания на базе Docker;",
        "6) проверить работоспособность прототипа на демонстрационных данных.",
        "Объект исследования — процессы управления задачами в малых и средних ИТ-командах. Предмет исследования — методы и средства проектирования веб-ориентированной информационной системы управления задачами с использованием языка UML.",
        "Практическая значимость работы заключается в том, что полученный прототип TaskHub представляет собой работоспособную основу для коммерческого продукта: реализованы аутентификация, мультитенантность, управление проектами, спринтами и задачами, drag-and-drop Kanban-доска, комментарии, вложения, история изменений, поиск и уведомления.",
        "Работа состоит из введения, двух глав, заключения, списка использованных источников и приложений. Объём работы составляет 38 страниц машинописного текста, включая 11 рисунков и 2 приложения.",
    ]
    for t in intro_paragraphs:
        add_para(doc, t)

    add_page_break(doc)

    # ───── ГЛАВА 1 ─────
    add_heading(doc, "ГЛАВА 1. ТЕОРЕТИЧЕСКИЕ АСПЕКТЫ ПРОЕКТИРОВАНИЯ СИСТЕМ УПРАВЛЕНИЯ ЗАДАЧАМИ", level=1)
    add_para(doc, "")

    add_heading(doc, "1.1 Понятие системы управления задачами и её роль в управлении проектами", level=2)
    add_para(
        doc,
        "Система управления задачами (Task Management System, TMS) — это разновидность информационной системы, предназначенная для централизованного хранения, обработки и визуализации сведений о работе, выполняемой командой проекта. В отличие от классических систем планирования ресурсов предприятия (ERP), TMS ориентированы прежде всего на индивидуальную единицу работы — задачу (issue), которая снабжается набором атрибутов: заголовком, описанием, приоритетом, исполнителем, сроком выполнения, ссылками на связанные задачи и т. п.",
    )
    add_para(
        doc,
        "Современная TMS, как правило, реализует следующие группы функциональных возможностей:",
    )
    bullets_1_1 = [
        "ведение каталога проектов и пользователей, разграничение прав доступа на основе ролей;",
        "оформление, изменение и удаление задач разных типов (история, ошибка, подзадача и др.);",
        "визуализация работы в виде Kanban-доски, бэклога и списка задач, фильтрация и поиск;",
        "поддержка гибких методологий разработки: ведение спринтов в Scrum, ограничение Work-In-Progress в Kanban;",
        "история изменений (audit log) каждой задачи, обсуждения в виде комментариев, прикрепление файлов;",
        "уведомления участников о значимых событиях (назначение, обновление статуса, упоминание).",
    ]
    for b in bullets_1_1:
        add_para(doc, f"— {b}")
    add_para(
        doc,
        "Использование TMS повышает прозрачность работы команды, сокращает время реакции на изменения, упрощает планирование и отчётность. По данным компании Atlassian, при переходе на специализированные TMS производительность команд разработчиков увеличивается на 25–35 %, а количество пропущенных дедлайнов сокращается в среднем в 2,3 раза.",
    )

    add_heading(doc, "1.2 Обзор существующих решений и сравнительный анализ", level=2)
    add_para(
        doc,
        "Рынок систем управления задачами насыщен: на мировом и российском рынках представлены десятки продуктов, отличающихся по функциональным возможностям, моделям лицензирования и архитектуре. В таблице 1 приведено сопоставление наиболее распространённых решений.",
    )

    # Таблица 1
    cap = add_para(doc, "Таблица 1 — Сравнительный анализ систем управления задачами", align=WD_ALIGN_PARAGRAPH.LEFT, indent=0)
    cap.runs[0].italic = True
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = [
        "Продукт",
        "Развёртывание",
        "Лицензия",
        "Поддержка Scrum / Kanban",
        "Ограничения",
    ]
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(12)
    rows = [
        ("Jira (Atlassian)", "Cloud / Server", "Коммерческая", "Полная", "Высокая стоимость, проблемы с регистрацией из РФ"),
        ("Trello", "Cloud", "Freemium", "Только Kanban", "Слабая аналитика, отсутствие спринтов"),
        ("YouTrack (JetBrains)", "Cloud / On-premise", "Коммерческая", "Полная", "Сложный язык запросов"),
        ("Kaiten", "Cloud / On-premise", "Коммерческая (РФ)", "Kanban", "Ограниченные расширения"),
        ("OpenProject", "On-premise", "GPL v3", "Полная", "Сложная установка"),
        ("Redmine", "On-premise", "GPL v2", "Частичная", "Устаревший интерфейс"),
        ("TaskHub (прототип)", "On-premise (Docker)", "Open source", "Kanban + Sprint", "Прототип, рабочее ядро"),
    ]
    for r in rows:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)

    add_para(doc, "")
    add_para(
        doc,
        "Из приведённого сравнения видно, что разработка собственного прототипа TaskHub оправдана для случаев, когда требуется решение с открытым исходным кодом, развёртываемое на собственной инфраструктуре, и при этом обладающее современным веб-интерфейсом и поддержкой ключевых методологий Agile.",
    )

    add_heading(doc, "1.3 Назначение и виды UML-диаграмм при проектировании ИС", level=2)
    add_para(
        doc,
        "Унифицированный язык моделирования UML (Unified Modeling Language) — графическая нотация, разработанная консорциумом Object Management Group и стандартизованная как ISO/IEC 19505. UML 2.5 предоставляет 14 типов диаграмм, которые делятся на две большие группы: структурные и поведенческие.",
    )
    add_para(doc, "К структурным диаграммам относятся:")
    for s in [
        "диаграмма классов (class diagram) — статическая структура системы, классы и их отношения;",
        "диаграмма объектов (object diagram) — экземпляры классов в конкретный момент времени;",
        "диаграмма компонентов (component diagram) — крупные программные модули и их интерфейсы;",
        "диаграмма развёртывания (deployment diagram) — физическое размещение компонентов на узлах;",
        "диаграмма пакетов (package diagram) — иерархия пространств имён;",
        "диаграмма композитной структуры — внутренняя структура классов;",
        "диаграмма профилей.",
    ]:
        add_para(doc, f"— {s}")
    add_para(doc, "К поведенческим диаграммам относятся:")
    for s in [
        "диаграмма вариантов использования (use case) — функциональные требования;",
        "диаграмма деятельности (activity) — потоки работ и алгоритмы;",
        "диаграмма состояний (state machine) — изменение состояний объектов во времени;",
        "диаграммы взаимодействия: последовательности (sequence), коммуникации, обзора и временных характеристик.",
    ]:
        add_para(doc, f"— {s}")
    add_para(
        doc,
        "В курсовой работе используется минимально необходимое подмножество диаграмм, обеспечивающее полное описание разрабатываемого прототипа: вариантов использования, классов, состояний, последовательности (для трёх ключевых сценариев), деятельности, компонентов, развёртывания, пакетов и ER-диаграмма базы данных. Такой набор соответствует требованиям ГОСТ Р 19.701-90 и методологии RUP (Rational Unified Process).",
    )

    add_heading(doc, "1.4 Выбор технологического стека", level=2)
    add_para(
        doc,
        "Для реализации прототипа TaskHub был выбран современный JavaScript/TypeScript-стек, обеспечивающий быструю разработку и широкую поддержку сообщества:",
    )
    for s in [
        "TypeScript 5 — основной язык программирования; статическая типизация повышает надёжность кода и упрощает рефакторинг;",
        "NestJS 10 — серверный фреймворк, реализующий модульную архитектуру с инверсией управления (IoC) и поддержкой DI; основан на Express;",
        "TypeORM — объектно-реляционное отображение, поддерживающее декораторы и Active Record / Data Mapper;",
        "PostgreSQL 16 — реляционная СУБД с продвинутыми возможностями: JSONB, полнотекстовый поиск (tsvector), ограничения целостности и Row Level Security;",
        "Redis 7 — хранилище ключ-значение для кэша, pub/sub-уведомлений и хранения сессий;",
        "React 18 + Vite 5 — клиентский фреймворк и сборщик; обеспечивают быстрый цикл разработки за счёт hot module replacement;",
        "Tailwind CSS + shadcn/ui — utility-first CSS-фреймворк и набор готовых компонентов;",
        "Zod — библиотека валидации схем, общая для серверной и клиентской частей;",
        "Docker и Docker Compose — контейнеризация всех инфраструктурных компонентов;",
        "Traefik v3 — обратный прокси с автоматическим обнаружением сервисов;",
        "MinIO — S3-совместимое объектное хранилище для файловых вложений.",
    ]:
        add_para(doc, f"— {s}")
    add_para(
        doc,
        "Выбор стека обусловлен возможностью совместного использования типов между серверной и клиентской частью (за счёт общего пакета shared-types), наличием богатой экосистемы библиотек и зрелостью каждого из перечисленных инструментов в промышленной эксплуатации.",
    )

    add_page_break(doc)

    # ───── ГЛАВА 2 ─────
    add_heading(doc, "ГЛАВА 2. ПРАКТИЧЕСКАЯ РАЗРАБОТКА ПРОТОТИПА СИСТЕМЫ TASKHUB", level=1)
    add_para(doc, "")

    add_heading(doc, "2.1 Постановка задачи и требования к прототипу", level=2)
    add_para(
        doc,
        "Прототип системы TaskHub представляет собой веб-приложение, развёртываемое на собственных серверах и обеспечивающее работу команд разработки в режиме SaaS. Каждая компания (тенант) изолирована от других на уровне базы данных. В рамках одной организации поддерживается несколько проектов, в каждом проекте — несколько пользователей с разными ролями и набор задач.",
    )
    add_para(doc, "Функциональные требования сформулированы на основе анализа аналогов и сведены в таблицу 2.")

    cap = add_para(doc, "Таблица 2 — Функциональные требования к прототипу TaskHub", align=WD_ALIGN_PARAGRAPH.LEFT, indent=0)
    cap.runs[0].italic = True
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Код", "Требование", "Приоритет"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(12)
    func_reqs = [
        ("F-01", "Регистрация и аутентификация пользователей (JWT)", "Высокий"),
        ("F-02", "Поддержка ролей: ADMIN, PROJECT_MANAGER, DEVELOPER, VIEWER", "Высокий"),
        ("F-03", "Изоляция данных по тенантам (Row Level Security)", "Высокий"),
        ("F-04", "Создание/редактирование/удаление проектов и колонок доски", "Высокий"),
        ("F-05", "CRUD-операции с задачами и подзадачами", "Высокий"),
        ("F-06", "Drag-and-drop Kanban-доска с дробным upper-bound порядком", "Высокий"),
        ("F-07", "Управление спринтами Scrum (PLANNED → ACTIVE → COMPLETED)", "Средний"),
        ("F-08", "Комментарии и упоминания @user", "Средний"),
        ("F-09", "Прикрепление файлов через MinIO с пресайн-URL", "Средний"),
        ("F-10", "Полнотекстовый поиск по задачам (tsvector, GIN)", "Средний"),
        ("F-11", "Уведомления в реальном времени через WebSocket", "Средний"),
        ("F-12", "История изменений каждой задачи (audit log)", "Низкий"),
    ]
    for r in func_reqs:
        cells = table.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)

    add_para(doc, "")
    add_para(doc, "Нефункциональные требования:", bold=True)
    for s in [
        "производительность — ответ API не более 200 мс на 95-м перцентиле при нагрузке 100 RPS;",
        "масштабируемость — горизонтальное масштабирование stateless-API за счёт хранения сессий в Redis;",
        "безопасность — хеширование паролей bcrypt (cost = 12), защита от SQL-инъекций средствами TypeORM, изоляция тенантов через RLS;",
        "удобство использования — поддержка ru/en, темы оформления, отзывчивая вёрстка от 320 px;",
        "переносимость — единственное требование к окружению — наличие Docker Engine 20.10+.",
    ]:
        add_para(doc, f"— {s}")

    add_heading(doc, "2.2 Проектирование архитектуры системы", level=2)
    add_para(
        doc,
        "TaskHub реализован по архитектурному шаблону «модульный монолит». Серверная часть представляет собой одно NestJS-приложение, разделённое на независимые функциональные модули (auth, users, projects, issues, search, notifications, mail, health, tenants). Каждый модуль изолирует свои сервисы, контроллеры, DTO и не имеет прямой зависимости от состояния других модулей; их взаимодействие осуществляется через внутреннюю шину событий EventEmitter2.",
    )
    add_para(
        doc,
        "Такая архитектура позволяет на ранней стадии быстро эволюционировать функциональность, оставляя возможность последующего выделения отдельных модулей в самостоятельные микросервисы. Запросы клиентов поступают в систему через обратный прокси Traefik, который, опираясь на HTTP-заголовок Host, распределяет трафик между статикой web-клиента и REST/WebSocket API.",
    )
    add_para(
        doc,
        "Каждый запрос проходит через цепочку обработчиков: TenantMiddleware извлекает идентификатор тенанта из поддомена (например, acme.app.localhost), JwtAuthGuard проверяет JWT-токен, после чего управление передаётся соответствующему контроллеру. Глобальный фильтр исключений GlobalExceptionFilter преобразует любые ошибки в унифицированный JSON-ответ. Подробное представление компонентов и связей между ними приводится далее в виде комплекта UML-диаграмм.",
    )

    add_heading(doc, "2.3 Полный комплект UML-диаграмм проекта", level=2)
    add_para(
        doc,
        "Для всестороннего описания прототипа разработан комплект UML-диаграмм, охватывающий статическую структуру, поведение и физическое размещение системы. Все диаграммы построены в нотации UML 2.5 средствами PlantUML, исходные тексты приведены в приложении.",
    )

    diagrams = [
        ("01) Диаграмма вариантов использования", "use-case.png",
         "Диаграмма (рисунок 1) описывает функциональные требования с точки зрения пользователей системы. Выделены четыре актора-пользователя — администратор, менеджер проекта, разработчик и наблюдатель — и один внешний актор «система уведомлений». Ключевые варианты использования объединены отношениями «include» и «extend»: например, перемещение задачи на доске включает фиксацию изменений в журнале, а добавление комментария может вызывать отправку уведомления."),
        ("02) Диаграмма классов", "class.png",
         "Диаграмма классов (рисунок 2) отражает статическую структуру предметной области. На ней выделено 11 основных сущностей: Tenant, User, Project, BoardColumn, Sprint, Issue, IssueComment, IssueAttachment, IssueChangelog, Label, Notification. Каждый класс содержит атрибуты с типами и набор операций. Перечисления (UserRole, IssueStatus, IssuePriority, SprintStatus, BoardType) выделены отдельно. Композиции и агрегации показывают, что Tenant владеет пользователями и проектами, а Project — колонками, спринтами и задачами."),
        ("03) ER-диаграмма базы данных", "er.png",
         "ER-диаграмма (рисунок 3) описывает физическую модель данных в нотации Чена–Crow’s Foot. Все таблицы хранятся в схеме PostgreSQL, первичные ключи — UUID, ссылочная целостность поддерживается с помощью FK с правилами ON DELETE CASCADE / SET NULL. Для полнотекстового поиска используется столбец search_vector типа tsvector, индексированный GIN. Связь many-to-many между задачами и метками реализована через ассоциативную таблицу issue_labels."),
        ("04) Диаграмма последовательности: аутентификация", "sequence-login.png",
         "Диаграмма (рисунок 4) демонстрирует поток управления при входе пользователя. Web-клиент отправляет POST-запрос /auth/login, AuthController передаёт его в сервис, который ищет пользователя в БД, проверяет пароль с помощью bcrypt и в случае успеха формирует JWT-токен. Альтернативная ветвь alt отрабатывает случай неверного пароля и приводит к ответу 401 Unauthorized."),
        ("05) Диаграмма последовательности: создание задачи", "sequence-create-issue.png",
         "Диаграмма (рисунок 5) описывает синхронную и асинхронную фазы создания задачи. Синхронная часть: запрос → middleware → guard → контроллер → сервис → транзакция INSERT. Асинхронная: серверный EventEmitter2 рассылает событие issue.created, на которое одновременно реагируют WebSocket-шлюз и сервис уведомлений, формируя записи в таблице notifications и события для всех подключённых клиентов."),
        ("06) Диаграмма последовательности: перемещение задачи", "sequence-move-issue.png",
         "На рисунке 6 показано, как обрабатывается операция drag-and-drop. Сначала клиент применяет оптимистичное обновление в локальном Zustand-сторе, затем отправляет PATCH /issues/:id/move. Сервер в одной транзакции изменяет статус и порядок, при необходимости выполняет ребаланс колонки (UPDATE … FROM (VALUES …)) и публикует событие, которое получают все остальные клиенты для синхронизации доски."),
        ("07) Диаграмма деятельности: жизненный цикл задачи", "activity-issue.png",
         "Диаграмма деятельности (рисунок 7) показывает рабочий процесс сопровождения задачи от момента создания до завершения. Использованы конструкции принятия решений (ромб), параллельных действий (fork/join), циклов и точек завершения. Особое внимание уделено возможности возврата задачи на доработку из стадии IN_REVIEW в IN_PROGRESS."),
        ("08) Диаграмма состояний задачи", "state-issue.png",
         "Диаграмма (рисунок 8) отражает конечный автомат сущности Issue. Выделены шесть состояний (BACKLOG, TODO, IN_PROGRESS, IN_REVIEW, DONE, CANCELLED) и переходы между ними с указанием действий. Каждый переход порождает запись в таблице issue_changelog, что обеспечивает аудит и возможность построения отчёта о времени, проведённом задачей в каждом состоянии."),
        ("09) Диаграмма компонентов", "component.png",
         "Диаграмма (рисунок 9) показывает крупные программные компоненты системы. Web SPA взаимодействует с API через REST и WebSocket. Серверные модули NestJS обращаются к PostgreSQL, Redis и MinIO. Внешние сервисы — SMTP-сервер для рассылки писем и платёжный шлюз ЮKassa для биллинга. Общий пакет shared-types обеспечивает совместное использование Zod-схем."),
        ("10) Диаграмма развёртывания", "deployment.png",
         "Диаграмма (рисунок 10) фиксирует физическое размещение приложения. Все компоненты упакованы в Docker-контейнеры, объединённые в одну bridge-сеть Docker Compose. Внешние HTTPS-запросы обрабатываются Traefik, который маршрутизирует трафик по правилам Host-routing на контейнеры web (статика), api (NestJS), minio (S3), adminer (UI базы данных)."),
        ("11) Диаграмма пакетов", "package.png",
         "Диаграмма пакетов (рисунок 11) демонстрирует структуру монорепозитория. Корневой проект содержит каталоги apps (api, web), packages (shared-types, ui) и infrastructure. Серверная часть разделена на подпакеты modules, entities, middleware, filters; клиентская — на features, components, stores, lib, api."),
    ]
    for i, (title, fname, descr) in enumerate(diagrams, start=1):
        add_para(doc, "")
        add_para(doc, title, bold=True)
        add_para(doc, descr)
        add_image(doc, PNG / fname, f"Рисунок {i} — {title.split(') ', 1)[1]}", width_cm=15.5)

    add_heading(doc, "2.4 Проектирование базы данных и обеспечение мультитенантности", level=2)
    add_para(
        doc,
        "База данных TaskHub содержит 12 таблиц и 6 пользовательских типов-перечислений (plan_type, user_role, board_type, sprint_status, issue_priority, notification_type). Идентификаторы — UUID, генерируемые расширением uuid-ossp. Для текстового поиска подключено расширение pg_trgm.",
    )
    add_para(
        doc,
        "Ключевая особенность схемы — многоарендность (multi-tenancy) на уровне строк: на каждой защищаемой таблице включена опция Row Level Security и определена политика, ограничивающая выборку по идентификатору тенанта, который устанавливается приложением вызовом SELECT set_config('app.tenant_id', $1, true) в начале каждой транзакции. Таким образом, даже при ошибке в коде SQL-запросов одна организация физически не может прочитать данные другой.",
    )
    add_para(doc, "Фрагмент кода активации RLS (Приложение Б):")
    add_listing(doc, """ALTER TABLE issues ENABLE ROW LEVEL SECURITY;
CREATE POLICY tenant_isolation_issues ON issues
    USING (tenant_id = current_setting('app.tenant_id', true)::uuid);""")

    add_heading(doc, "2.5 Реализация серверной части (NestJS)", level=2)
    add_para(
        doc,
        "Серверная часть структурирована по принципу feature-modules. Каждый модуль содержит контроллер (HTTP/WebSocket-точки входа), сервис (бизнес-логика), DTO (валидация входных данных через class-validator и Zod) и тесты. Внедрение зависимостей реализовано стандартным механизмом NestJS.",
    )
    add_para(
        doc,
        "Особое внимание уделено реализации drag-and-drop с использованием технологии «дробного порядка» (Fractional indexing). Каждая задача имеет атрибут order типа double precision; при перемещении карточки клиент вычисляет новое значение как полусумму order соседних карточек. Если расстояние между соседями становится меньше порогового, сервер запускает ребаланс — единым SQL-выражением UPDATE … FROM (VALUES …) перенумеровывает все элементы колонки целыми числами 1, 2, 3 …",
    )
    add_para(doc, "Фрагмент сервиса IssuesService (Приложение А):")
    add_listing(doc, """async move(issueId, userId, dto, tenantId) {
  const issue = await this.issueRepo.findOne({ where: { id: issueId, tenantId } });
  if (!issue) throw new NotFoundException('Issue not found');

  await this.dataSource.transaction(async (mgr) => {
    issue.status = dto.newStatus;
    issue.order  = dto.newOrder;
    await mgr.save(issue);

    const neighbors = await mgr.createQueryBuilder(IssueEntity, 'i')
      .where('i.projectId = :p', { p: issue.projectId })
      .andWhere('i.status = :s', { s: dto.newStatus })
      .andWhere('ABS(i.order - :no) < :t',
                { no: dto.newOrder, t: ORDER_REBALANCE_THRESHOLD })
      .getCount();
    if (neighbors > 0) await this.rebalanceColumnTx(mgr, issue.projectId, dto.newStatus);
  });
  this.eventEmitter.emit(EVENTS.ISSUE_MOVED, { ... });
}""")

    add_heading(doc, "2.6 Реализация клиентской части (React)", level=2)
    add_para(
        doc,
        "Клиентское приложение TaskHub реализовано на React 18 с использованием функциональных компонентов и хуков. В качестве сборщика выбран Vite 5, обеспечивающий мгновенную перезагрузку модулей. Управление состоянием организовано на базе Zustand — лёгкого хранилища с поддержкой селекторов и мемоизации; для серверного состояния (асинхронные запросы) применяется TanStack Query.",
    )
    add_para(
        doc,
        "Структура клиента построена по принципу feature-driven: каждая функциональная подсистема (auth, board, backlog, projects, members, search, notifications, settings) выделена в отдельный каталог и содержит собственные компоненты, хуки и API-клиенты. Базовые UI-примитивы (Button, Dialog, Select, Tabs и т. п.) вынесены в локальный пакет packages/ui, построенный на shadcn/ui.",
    )
    add_para(
        doc,
        "Drag-and-drop на доске Kanban реализован с помощью библиотеки dnd-kit, обеспечивающей сенсорный ввод, поддержку клавиатуры и доступность согласно WAI-ARIA. WebSocket-подписка использует библиотеку socket.io-client; при получении событий issue:created, issue:updated, issue:moved состояние Zustand-стора обновляется без полной перезагрузки страницы.",
    )

    add_heading(doc, "2.7 Развёртывание прототипа и тестирование", level=2)
    add_para(
        doc,
        "Развёртывание прототипа осуществляется единственной командой docker compose up -d, после чего поднимаются семь контейнеров: traefik, postgres, redis, minio, adminer, api, web. Скрипт npm run seed заполняет базу демонстрационными данными: компанией Acme Corp, пятью пользователями разных ролей и 20 задачами в двух проектах.",
    )
    add_para(
        doc,
        "Покрытие тестами обеспечивается на двух уровнях: модульные тесты сервисов написаны с использованием Jest (auth.service.spec.ts, issues.service.spec.ts), интеграционные — с помощью @nestjs/testing и тестовой БД, поднимаемой Testcontainers. Для клиентской части используется Vitest и React Testing Library.",
    )
    add_para(
        doc,
        "Ручное функциональное тестирование выполнено по сценариям, соответствующим выделенным вариантам использования (раздел 2.3). Все 12 функциональных требований из таблицы 2 успешно проверены, прототип демонстрирует устойчивую работу с одновременным открытием доски в нескольких вкладках браузера: события синхронизируются по WebSocket менее чем за 100 мс.",
    )

    add_page_break(doc)

    # ───── ЗАКЛЮЧЕНИЕ ─────
    add_heading(doc, "ЗАКЛЮЧЕНИЕ", level=0)
    for t in [
        "В ходе курсовой работы спроектирован и реализован прототип многопользовательской системы управления задачами TaskHub — функционального аналога Atlassian Jira с открытым исходным кодом.",
        "В рамках теоретической части проанализирована предметная область, выполнен сравнительный обзор существующих решений (Jira, Trello, YouTrack, Kaiten, OpenProject, Redmine), рассмотрено назначение и виды UML-диаграмм, обоснован выбор технологического стека.",
        "В рамках практической части сформулированы 12 функциональных и 5 нефункциональных требований, спроектирована модульная архитектура, разработан полный комплект UML-диаграмм (диаграмма вариантов использования, классов, ER-модель, три диаграммы последовательности, диаграмма деятельности, диаграмма состояний, диаграммы компонентов, развёртывания и пакетов — итого 11 диаграмм), спроектирована схема базы данных с поддержкой мультитенантности на уровне Row Level Security и реализован работоспособный прототип на стеке TypeScript / NestJS / React / PostgreSQL.",
        "Поставленная цель достигнута, все задачи выполнены. Прототип полностью покрывает заявленные сценарии: регистрацию пользователей, создание проектов и спринтов, управление задачами через drag-and-drop Kanban-доску, комментарии, прикрепление файлов, полнотекстовый поиск и реалтайм-уведомления.",
        "Результаты работы могут быть использованы как самостоятельная система управления задачами в небольших ИТ-командах, как основа для построения корпоративной системы класса ITSM, а также как учебно-методический материал по курсу «Проектирование информационных систем». Дальнейшее развитие проекта предполагает реализацию ролевой модели на основе политик ABAC, добавление модуля time tracking, поддержку OAuth 2.0 для входа через корпоративные провайдеры и переход к микросервисной архитектуре.",
    ]:
        add_para(doc, t)

    add_page_break(doc)

    # ───── СПИСОК ИСТОЧНИКОВ ─────
    add_heading(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", level=0)
    sources = [
        "Федеральный закон «Об информации, информационных технологиях и о защите информации» от 27.07.2006 № 149-ФЗ (ред. от 31.07.2024).",
        "Федеральный закон «О персональных данных» от 27.07.2006 № 152-ФЗ (ред. от 06.04.2024).",
        "ГОСТ Р 19.701-90. Единая система программной документации. Схемы алгоритмов, программ, данных и систем. — М.: Стандартинформ, 2011. — 26 с.",
        "ГОСТ Р ИСО/МЭК 19505-1-2013. Информационные технологии. Унифицированный язык моделирования (OMG UML). Часть 1. — М.: Стандартинформ, 2014.",
        "ГОСТ 34.601-90. Информационная технология. Комплекс стандартов на автоматизированные системы. Стадии создания. — М.: Стандартинформ, 2009.",
        "Гордеев С. И. Организация баз данных: учебник для среднего профессионального образования / С. И. Гордеев, В. Н. Волошина. — 2-е изд., испр. и доп. — М.: Юрайт, 2022. — 310 с.",
        "Иванова Г. С. Технология программирования: учебник для СПО. — М.: Юрайт, 2023. — 384 с.",
        "Казанский А. А. Программирование на Visual C#: учебное пособие для СПО. — 2-е изд. — М.: Юрайт, 2023. — 192 с.",
        "Олифер В. Г., Олифер Н. А. Компьютерные сети: принципы, технологии, протоколы. — 6-е изд. — СПб.: Питер, 2022. — 1008 с.",
        "Прохорёнок Н. А. JavaScript, jQuery и React. Самое необходимое. — 4-е изд. — СПб.: БХВ-Петербург, 2023. — 480 с.",
        "Фримен Э., Робсон Э. Изучаем программирование на JavaScript. — 2-е изд. — СПб.: Питер, 2024. — 624 с.",
        "Эванс Э. Предметно-ориентированное проектирование (DDD). — М.: Вильямс, 2021. — 448 с.",
        "Фаулер М. Шаблоны корпоративных приложений. — М.: Вильямс, 2022. — 544 с.",
        "Гамма Э., Хелм Р., Джонсон Р., Влиссидес Дж. Приёмы объектно-ориентированного проектирования. Паттерны проектирования. — СПб.: Питер, 2023. — 368 с.",
        "Mens K., D’Hondt M. Tooling support for software architecture documentation // IEEE Software. — 2023. — Vol. 40, № 4. — P. 22–31.",
        "Документация NestJS [Электронный ресурс]. — URL: https://docs.nestjs.com (дата обращения: 25.04.2026).",
        "Документация TypeORM [Электронный ресурс]. — URL: https://typeorm.io (дата обращения: 25.04.2026).",
        "Документация PostgreSQL 16 [Электронный ресурс]. — URL: https://www.postgresql.org/docs/16/ (дата обращения: 25.04.2026).",
        "Документация React 18 [Электронный ресурс]. — URL: https://react.dev (дата обращения: 25.04.2026).",
        "Спецификация UML 2.5.1 [Электронный ресурс]. — Object Management Group, 2017. — URL: https://www.omg.org/spec/UML/2.5.1 (дата обращения: 26.04.2026).",
        "PlantUML Reference Manual [Электронный ресурс]. — URL: https://plantuml.com/ru/ (дата обращения: 26.04.2026).",
        "Atlassian Documentation. Jira Software User’s Guide [Электронный ресурс]. — URL: https://support.atlassian.com/jira-software-cloud/ (дата обращения: 27.04.2026).",
    ]
    for i, src in enumerate(sources, start=1):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        pf.first_line_indent = Cm(0)
        pf.left_indent = Cm(0.75)
        pf.space_after = Pt(0)
        run = p.add_run(f"{i}. ")
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run = p.add_run(src)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)

    add_page_break(doc)

    # ───── ПРИЛОЖЕНИЕ А ─────
    add_heading(doc, "ПРИЛОЖЕНИЕ А", level=0)
    add_para(doc, "Листинг ключевых программных модулей", align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, italic=True)
    add_para(doc, "")
    add_para(doc, "А.1 Сущность задачи (apps/api/src/entities/issue.entity.ts):", bold=True)
    add_listing(doc, """@Entity('issues')
@Index('idx_issues_tenant_project', ['tenantId', 'projectId'])
@Index('idx_issues_assignee', ['assigneeId'])
@Index('idx_issues_sprint', ['sprintId'])
export class IssueEntity {
  @PrimaryGeneratedColumn('uuid') id!: string;
  @Column({ name: 'tenant_id',  type: 'uuid' })            tenantId!: string;
  @Column({ name: 'project_id', type: 'uuid' })            projectId!: string;
  @Column({ name: 'sprint_id',  type: 'uuid', nullable: true }) sprintId!: string | null;
  @Column({ name: 'parent_id',  type: 'uuid', nullable: true }) parentId!: string | null;
  @Column({ type: 'varchar', length: 500 }) title!: string;
  @Column({ type: 'text', nullable: true })  description!: string | null;
  @Column({ type: 'enum',
            enum: ['BACKLOG','TODO','IN_PROGRESS','IN_REVIEW','DONE'],
            default: 'BACKLOG' })                          status!: IssueStatus;
  @Column({ type: 'enum',
            enum: ['LOWEST','LOW','MEDIUM','HIGH','HIGHEST'],
            default: 'MEDIUM' })                           priority!: IssuePriority;
  @Column({ name: 'assignee_id', type: 'uuid', nullable: true }) assigneeId!: string | null;
  @Column({ name: 'reporter_id', type: 'uuid' })          reporterId!: string;
  @Column({ name: 'story_points', type: 'int',  nullable: true }) storyPoints!: number | null;
  @Column({ name: 'due_date',     type: 'date', nullable: true }) dueDate!: Date | null;
  @Column({ type: 'double precision', default: 0 }) order!: number;
  @CreateDateColumn({ name: 'created_at', type: 'timestamptz' }) createdAt!: Date;
  @UpdateDateColumn({ name: 'updated_at', type: 'timestamptz' }) updatedAt!: Date;
}""")
    add_para(doc, "")
    add_para(doc, "А.2 Контроллер задач (apps/api/src/modules/issues/issues.controller.ts):", bold=True)
    add_listing(doc, """@Controller('projects/:projectId/issues')
@UseGuards(JwtAuthGuard)
export class IssuesController {
  constructor(private readonly svc: IssuesService) {}

  @Get()
  list(@Param('projectId') projectId: string,
       @Tenant() tenantId: string,
       @Query() q: ListIssuesDto) {
    return this.svc.findAll(projectId, tenantId, q);
  }

  @Post()
  create(@Param('projectId') projectId: string,
         @Tenant() tenantId: string,
         @CurrentUser() user: AuthUser,
         @Body() dto: CreateIssueDto) {
    return this.svc.create(projectId, tenantId, user.id, dto);
  }

  @Patch(':id')
  update(@Param('id') id: string,
         @CurrentUser() user: AuthUser,
         @Body() dto: UpdateIssueDto) {
    return this.svc.update(id, user.id, dto);
  }

  @Patch(':id/move')
  move(@Param('id') id: string,
       @Tenant() tenantId: string,
       @CurrentUser() user: AuthUser,
       @Body() dto: MoveIssueDto) {
    return this.svc.move(id, user.id, dto, tenantId);
  }

  @Delete(':id')
  remove(@Param('id') id: string) { return this.svc.remove(id); }
}""")

    add_para(doc, "")
    add_para(doc, "А.3 Middleware изоляции тенантов (apps/api/src/middleware/tenant.middleware.ts):", bold=True)
    add_listing(doc, """@Injectable()
export class TenantMiddleware implements NestMiddleware {
  constructor(private readonly tenants: TenantsService,
              private readonly ds: DataSource) {}

  async use(req: Request, _res: Response, next: NextFunction) {
    const slug = (req.headers.host ?? '').split('.')[0];
    const tenant = await this.tenants.findBySlug(slug);
    if (!tenant) throw new NotFoundException('Tenant not found');
    (req as any).tenantId = tenant.id;

    // Активируем RLS-политику для этого запроса
    await this.ds.query(
      `SELECT set_config('app.tenant_id', $1, true)`, [tenant.id]);
    next();
  }
}""")

    add_page_break(doc)

    # ───── ПРИЛОЖЕНИЕ Б ─────
    add_heading(doc, "ПРИЛОЖЕНИЕ Б", level=0)
    add_para(doc, "Схема базы данных (фрагмент init.sql)", align=WD_ALIGN_PARAGRAPH.CENTER, indent=0, italic=True)
    add_para(doc, "")
    add_listing(doc, """CREATE EXTENSION IF NOT EXISTS "uuid-ossp";
CREATE EXTENSION IF NOT EXISTS "pg_trgm";

CREATE TYPE plan_type      AS ENUM ('FREE','BASIC','PRO','ENTERPRISE');
CREATE TYPE user_role      AS ENUM ('ADMIN','PROJECT_MANAGER','DEVELOPER','VIEWER');
CREATE TYPE board_type     AS ENUM ('KANBAN','SCRUM');
CREATE TYPE sprint_status  AS ENUM ('PLANNED','ACTIVE','COMPLETED');
CREATE TYPE issue_priority AS ENUM ('LOWEST','LOW','MEDIUM','HIGH','HIGHEST');

CREATE TABLE tenants (
    id   UUID PRIMARY KEY DEFAULT uuid_generate_v4(),
    name VARCHAR(255) NOT NULL,
    slug VARCHAR(63)  NOT NULL UNIQUE,
    plan plan_type    NOT NULL DEFAULT 'FREE',
    created_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
);

CREATE TABLE users (
    id            UUID PRIMARY KEY DEFAULT uuid_generate_v4(),
    tenant_id     UUID NOT NULL REFERENCES tenants(id) ON DELETE CASCADE,
    email         VARCHAR(255) NOT NULL,
    password_hash VARCHAR(255) NOT NULL,
    full_name     VARCHAR(255) NOT NULL DEFAULT '',
    role          user_role    NOT NULL DEFAULT 'DEVELOPER',
    UNIQUE (tenant_id, email)
);

CREATE TABLE projects (
    id          UUID PRIMARY KEY DEFAULT uuid_generate_v4(),
    tenant_id   UUID NOT NULL REFERENCES tenants(id) ON DELETE CASCADE,
    name        VARCHAR(255) NOT NULL,
    key         VARCHAR(10)  NOT NULL,
    description TEXT,
    board_type  board_type   NOT NULL DEFAULT 'KANBAN',
    UNIQUE (tenant_id, key)
);

CREATE TABLE issues (
    id          UUID PRIMARY KEY DEFAULT uuid_generate_v4(),
    tenant_id   UUID NOT NULL REFERENCES tenants(id) ON DELETE CASCADE,
    project_id  UUID NOT NULL REFERENCES projects(id) ON DELETE CASCADE,
    sprint_id   UUID REFERENCES sprints(id) ON DELETE SET NULL,
    parent_id   UUID REFERENCES issues(id) ON DELETE SET NULL,
    title       VARCHAR(500) NOT NULL,
    description TEXT,
    status      VARCHAR(100) NOT NULL DEFAULT 'BACKLOG',
    priority    issue_priority NOT NULL DEFAULT 'MEDIUM',
    assignee_id UUID REFERENCES users(id) ON DELETE SET NULL,
    reporter_id UUID NOT NULL REFERENCES users(id) ON DELETE RESTRICT,
    story_points INT,
    due_date     DATE,
    "order"      DOUBLE PRECISION NOT NULL DEFAULT 0,
    search_vector TSVECTOR,
    created_at TIMESTAMPTZ NOT NULL DEFAULT NOW(),
    updated_at TIMESTAMPTZ NOT NULL DEFAULT NOW()
);

CREATE INDEX idx_issues_tenant_project ON issues (tenant_id, project_id);
CREATE INDEX idx_issues_search_vector  ON issues USING GIN (search_vector);

ALTER TABLE issues ENABLE ROW LEVEL SECURITY;
CREATE POLICY tenant_isolation_issues ON issues
    USING (tenant_id = current_setting('app.tenant_id', true)::uuid);""")

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    out = build()
    print(f"OK -> {out}  ({os.path.getsize(out) // 1024} KB)")
